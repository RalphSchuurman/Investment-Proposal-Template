import pandas as pd
import yfinance as yf
from docx import Document
from datetime import datetime
from docx.shared import Inches


def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)


def create_doc(companyName, sector, industry, current_price,
                fiftyTwoWeek, targetMeanPrice,
               marketCap, beta, dividendRate, companyInfo, companyLogo,
               shortRatio, shortPercentage, dividendHistory,
               news, competition_df,plot,author):
    document = Document('Empty_koopvoorstel.docx')


    # replace the text in paragraphs
    for paragraph in document.paragraphs:
        paragraphtext = paragraph.text
        find_replace("[COMPANY_INFO]",companyInfo, paragraph)
        find_replace("[Short Ratio:]", shortRatio, paragraph)
        find_replace("[Short % of Shares Outstanding:]", shortPercentage, paragraph)
    # replace the text in tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_replace("Bedrijf:", companyName, paragraph )
                    find_replace("Sector:", sector, paragraph)
                    find_replace("Industrie:", industry, paragraph)
                    find_replace("Huidige Prijs:", current_price, paragraph)
                    find_replace("52 weken l/h:", fiftyTwoWeek, paragraph)
                    find_replace("Doel over 1 jaar:", targetMeanPrice, paragraph)
                    find_replace("Beurswaarde:", marketCap, paragraph)
                    find_replace("Beta:", beta, paragraph)
                    find_replace("Voorwaarts dividend:", dividendRate, paragraph)
                    find_replace("Datum:", "Datum: " + str(datetime.today()), paragraph)
                    find_replace("Auteur:", "Auteur: " + str(author), paragraph)

    # input table for competitor
    competitor_df = competition_df
    table = document.add_table(competitor_df.shape[0]+1, competitor_df.shape[1])
    for j in range(competitor_df.shape[-1]):
        table.cell(0, j).text = competitor_df.columns[j]

    for i in range(competitor_df.shape[0]):
        for j in range(competitor_df.shape[-1]):
            table.cell(i + 1, j).text = str(competitor_df.values[i, j])

    # replace new table with old table
    old_competitor_table = document.tables[2]
    newTable = document.tables[4]
    old_competitor_table._element.getparent().replace(old_competitor_table._element,newTable._element )

    ## Dividends table

    dividend_df = dividendHistory
    dividend_df = dividend_df.reset_index()
    table2 = document.add_table(dividend_df.shape[0] + 1, dividend_df.shape[1])
    for j in range(dividend_df.shape[-1]):
        table2.cell(0, j).text = dividend_df.columns[j]

    for i in range(dividend_df.shape[0]):
        for j in range(dividend_df.shape[-1]):
            table2.cell(i + 1, j).text = str(dividend_df.values[i, j])
    # replace new table with old table
    old_dividend_table = document.tables[3]
    newTable = document.tables[4]
    old_dividend_table._element.getparent().replace(old_dividend_table._element, newTable._element)

    ## add graph
    keyword = "Koersgrafiek"
    for paragraph in document.paragraphs:
        if keyword in paragraph.text:
            r = paragraph.add_run()
            r.add_picture(plot)

    return(document)

